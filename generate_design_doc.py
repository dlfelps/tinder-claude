#!/usr/bin/env python3
"""Generate the Tinder-Claude Design Document in DOCX format with diagrams."""

import io
import os
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


# ---------------------------------------------------------------------------
# Diagram helper functions
# ---------------------------------------------------------------------------

def _save_fig_to_bytes(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf


def _draw_box(ax, x, y, w, h, text, color="#4A90D9", text_color="white",
              fontsize=9, fontweight="bold", alpha=1.0):
    box = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.05",
        facecolor=color, edgecolor="#2C3E50",
        linewidth=1.2, alpha=alpha
    )
    ax.add_patch(box)
    ax.text(x + w / 2, y + h / 2, text,
            ha="center", va="center",
            fontsize=fontsize, fontweight=fontweight,
            color=text_color, wrap=True)


def _draw_arrow(ax, x1, y1, x2, y2, label="", color="#2C3E50"):
    ax.annotate(
        "",
        xy=(x2, y2), xytext=(x1, y1),
        arrowprops=dict(
            arrowstyle="-|>",
            color=color, lw=1.5,
            connectionstyle="arc3,rad=0"
        ),
    )
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(mx, my + 0.06, label,
                ha="center", va="bottom", fontsize=7,
                color=color, style="italic")


# ---------------------------------------------------------------------------
# Diagram 1: High-Level Architecture
# ---------------------------------------------------------------------------

def create_architecture_diagram() -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(-0.5, 10.5)
    ax.set_ylim(-0.5, 6.5)
    ax.axis("off")
    ax.set_title("High-Level System Architecture", fontsize=14,
                 fontweight="bold", pad=15)

    # Client
    _draw_box(ax, 0, 2.5, 1.8, 1.2, "HTTP\nClient", "#E67E22", "white", 10)

    # FastAPI Gateway
    _draw_box(ax, 2.8, 2.5, 2.0, 1.2, "FastAPI\nApplication\n(main.py)",
              "#2ECC71", "white", 9)

    # Routes layer
    _draw_box(ax, 5.8, 4.5, 1.6, 0.9, "Users\nRoutes", "#3498DB")
    _draw_box(ax, 5.8, 3.1, 1.6, 0.9, "Feed\nRoutes", "#3498DB")
    _draw_box(ax, 5.8, 1.7, 1.6, 0.9, "Swipe\nRoutes", "#3498DB")

    # Services
    _draw_box(ax, 8.2, 3.8, 1.8, 0.9, "Feed\nService", "#9B59B6")
    _draw_box(ax, 8.2, 2.4, 1.8, 0.9, "Swipe\nService", "#9B59B6")

    # Store
    _draw_box(ax, 8.2, 0.5, 1.8, 1.2, "InMemory\nStore\n(Singleton)",
              "#E74C3C", "white", 9)

    # Labels for layers
    ax.text(0.9, 5.8, "Client Layer", ha="center", fontsize=8,
            fontweight="bold", color="#7F8C8D")
    ax.text(3.8, 5.8, "API Layer", ha="center", fontsize=8,
            fontweight="bold", color="#7F8C8D")
    ax.text(6.6, 5.8, "Route Layer", ha="center", fontsize=8,
            fontweight="bold", color="#7F8C8D")
    ax.text(9.1, 5.8, "Service Layer", ha="center", fontsize=8,
            fontweight="bold", color="#7F8C8D")

    # Arrows
    _draw_arrow(ax, 1.8, 3.1, 2.8, 3.1, "HTTP")
    _draw_arrow(ax, 4.8, 3.5, 5.8, 4.9)
    _draw_arrow(ax, 4.8, 3.1, 5.8, 3.5)
    _draw_arrow(ax, 4.8, 2.8, 5.8, 2.2)

    _draw_arrow(ax, 7.4, 4.5, 8.2, 4.2)
    _draw_arrow(ax, 7.4, 3.2, 8.2, 3.0)
    _draw_arrow(ax, 7.4, 2.0, 8.2, 2.6)

    _draw_arrow(ax, 9.1, 2.4, 9.1, 1.7, "CRUD")
    _draw_arrow(ax, 9.1, 3.8, 9.1, 1.7, "Read")

    # Border for Data layer
    border = plt.Rectangle((7.9, 0.2), 2.4, 5.0, linewidth=1.5,
                           edgecolor="#BDC3C7", facecolor="none",
                           linestyle="--")
    ax.add_patch(border)

    return _save_fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# Diagram 2: Data Flow – User Request Lifecycle
# ---------------------------------------------------------------------------

def create_request_lifecycle_diagram() -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_xlim(-0.5, 11)
    ax.set_ylim(-0.5, 7.5)
    ax.axis("off")
    ax.set_title("User Request Lifecycle – Data Flow", fontsize=14,
                 fontweight="bold", pad=15)

    y_positions = {
        "create_user": 6.0,
        "get_feed": 4.2,
        "swipe": 2.4,
        "get_matches": 0.6,
    }

    colors = {
        "create_user": "#27AE60",
        "get_feed": "#2980B9",
        "swipe": "#8E44AD",
        "get_matches": "#D35400",
    }

    labels = {
        "create_user": "1. Create User  (POST /users/)",
        "get_feed": "2. Get Feed  (GET /feed?user_id=...)",
        "swipe": "3. Swipe  (POST /swipe)",
        "get_matches": "4. Get Matches  (GET /matches?user_id=...)",
    }

    for key, y in y_positions.items():
        c = colors[key]
        # Step label
        ax.text(0, y + 0.7, labels[key], fontsize=9.5, fontweight="bold",
                color=c)

        # Request box
        _draw_box(ax, 0, y - 0.25, 2.2, 0.65, "Request\nValidation\n(Pydantic)",
                  c, "white", 7.5)

        if key == "create_user":
            _draw_box(ax, 3.0, y - 0.25, 2.2, 0.65,
                      "Create User\nobject (UUID)", c, "white", 7.5)
            _draw_box(ax, 6.0, y - 0.25, 2.2, 0.65,
                      "Store\n.add_user()", c, "white", 7.5)
            _draw_box(ax, 9.0, y - 0.25, 1.8, 0.65,
                      "ApiResponse\n{data: User}", c, "white", 7.5)
        elif key == "get_feed":
            _draw_box(ax, 3.0, y - 0.25, 2.2, 0.65,
                      "FeedService\n.generate_feed()", c, "white", 7.5)
            _draw_box(ax, 6.0, y - 0.25, 2.2, 0.65,
                      "Filter:\nZone → Self → Seen", c, "white", 7.5)
            _draw_box(ax, 9.0, y - 0.25, 1.8, 0.65,
                      "ApiResponse\n{data: [Users]}", c, "white", 7.5)
        elif key == "swipe":
            _draw_box(ax, 3.0, y - 0.25, 2.2, 0.65,
                      "SwipeService\n.process_swipe()", c, "white", 7.5)
            _draw_box(ax, 6.0, y - 0.25, 2.2, 0.65,
                      "Check reverse\nLIKE → Match?", c, "white", 7.5)
            _draw_box(ax, 9.0, y - 0.25, 1.8, 0.65,
                      "ApiResponse\n{is_match: bool}", c, "white", 7.5)
        else:
            _draw_box(ax, 3.0, y - 0.25, 2.2, 0.65,
                      "Validate user\nexists", c, "white", 7.5)
            _draw_box(ax, 6.0, y - 0.25, 2.2, 0.65,
                      "Store\n.get_matches()", c, "white", 7.5)
            _draw_box(ax, 9.0, y - 0.25, 1.8, 0.65,
                      "ApiResponse\n{data: [Match]}", c, "white", 7.5)

        # Arrows between steps
        _draw_arrow(ax, 2.2, y + 0.08, 3.0, y + 0.08, color=c)
        _draw_arrow(ax, 5.2, y + 0.08, 6.0, y + 0.08, color=c)
        _draw_arrow(ax, 8.2, y + 0.08, 9.0, y + 0.08, color=c)

    return _save_fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# Diagram 3: Match Detection Dataflow
# ---------------------------------------------------------------------------

def create_match_detection_diagram() -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(-0.5, 10.5)
    ax.set_ylim(-1, 7.5)
    ax.axis("off")
    ax.set_title("Match Detection – Detailed Data Flow", fontsize=14,
                 fontweight="bold", pad=15)

    # Alice swipes Bob
    _draw_box(ax, 0, 5.5, 2.0, 1.0, "Alice\n(User A)", "#E67E22")
    _draw_arrow(ax, 2.0, 6.0, 3.2, 6.0, "LIKE →")
    _draw_box(ax, 3.2, 5.5, 2.0, 1.0, "POST /swipe\n{A→B, LIKE}", "#3498DB")
    _draw_arrow(ax, 5.2, 6.0, 6.4, 6.0)
    _draw_box(ax, 6.4, 5.5, 2.5, 1.0, "SwipeService\nprocess_swipe()", "#9B59B6")

    # Check reverse
    _draw_arrow(ax, 7.65, 5.5, 7.65, 4.5)
    _draw_box(ax, 6.0, 3.5, 3.2, 1.0, "Store.find_swipe(B→A)\nReverse swipe exists?",
              "#1ABC9C")

    # No reverse – store only
    _draw_arrow(ax, 6.0, 4.0, 4.5, 4.0, "No reverse")
    _draw_box(ax, 2.0, 3.5, 2.5, 1.0, "Store.add_swipe()\nis_match = False",
              "#E74C3C")

    # Bob swipes Alice
    _draw_box(ax, 0, 1.8, 2.0, 1.0, "Bob\n(User B)", "#E67E22")
    _draw_arrow(ax, 2.0, 2.3, 3.2, 2.3, "LIKE →")
    _draw_box(ax, 3.2, 1.8, 2.0, 1.0, "POST /swipe\n{B→A, LIKE}", "#3498DB")
    _draw_arrow(ax, 5.2, 2.3, 6.4, 2.3)
    _draw_box(ax, 6.4, 1.8, 2.5, 1.0, "SwipeService\nprocess_swipe()", "#9B59B6")

    # Check reverse – found
    _draw_arrow(ax, 7.65, 1.8, 7.65, 0.9)
    _draw_box(ax, 6.0, -0.1, 3.2, 1.0,
              "Store.find_swipe(A→B)\nFound! A LIKED B ✓", "#1ABC9C")

    # Match!
    _draw_arrow(ax, 6.0, 0.4, 4.5, 0.4, "Match found!")
    _draw_box(ax, 1.5, -0.1, 3.0, 1.0,
              "Store.add_match(A, B)\nis_match = True  ★", "#27AE60", "white", 9)

    return _save_fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# Diagram 4: Feed Generation – Filter Pipeline
# ---------------------------------------------------------------------------

def create_feed_pipeline_diagram() -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(-0.5, 10.5)
    ax.set_ylim(-0.5, 5)
    ax.axis("off")
    ax.set_title("Feed Generation – Three-Tier Filter Pipeline", fontsize=14,
                 fontweight="bold", pad=15)

    # All users pool
    _draw_box(ax, 0, 2, 1.8, 1.2, "All Users\nin Store\n(N users)", "#95A5A6")
    _draw_arrow(ax, 1.8, 2.6, 2.6, 2.6, "")

    # Filter 1: Zone
    _draw_box(ax, 2.6, 2, 2.0, 1.2, "Zone Filter\nzone_id ==\nrequester.zone_id",
              "#E74C3C")
    ax.text(3.6, 1.7, "Removes other zones", fontsize=6.5,
            ha="center", color="#E74C3C", style="italic")
    _draw_arrow(ax, 4.6, 2.6, 5.2, 2.6, "")

    # Filter 2: Self
    _draw_box(ax, 5.2, 2, 2.0, 1.2, "Self-Exclusion\nid !=\nrequester.id",
              "#E67E22")
    ax.text(6.2, 1.7, "Removes self", fontsize=6.5,
            ha="center", color="#E67E22", style="italic")
    _draw_arrow(ax, 7.2, 2.6, 7.8, 2.6, "")

    # Filter 3: Seen
    _draw_box(ax, 7.8, 2, 2.2, 1.2, "Seen-State\nFilter\nid ∉ swiped_ids",
              "#8E44AD")
    ax.text(8.9, 1.7, "Removes already swiped", fontsize=6.5,
            ha="center", color="#8E44AD", style="italic")

    # Output
    _draw_arrow(ax, 8.9, 3.2, 8.9, 3.8, "")
    _draw_box(ax, 7.8, 3.8, 2.2, 0.8, "Filtered\nFeed Result", "#27AE60")

    # Example counts
    ax.text(0.9, 3.5, "100", fontsize=12, ha="center",
            fontweight="bold", color="#95A5A6")
    ax.text(3.6, 3.5, "40", fontsize=12, ha="center",
            fontweight="bold", color="#E74C3C")
    ax.text(6.2, 3.5, "39", fontsize=12, ha="center",
            fontweight="bold", color="#E67E22")
    ax.text(8.9, 4.8, "25", fontsize=12, ha="center",
            fontweight="bold", color="#27AE60")

    return _save_fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# Diagram 5: Data Model Entity Relationship
# ---------------------------------------------------------------------------

def create_data_model_diagram() -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(-0.5, 10.5)
    ax.set_ylim(-0.5, 5.5)
    ax.axis("off")
    ax.set_title("Data Model – Entity Relationships", fontsize=14,
                 fontweight="bold", pad=15)

    # User entity
    user_box = FancyBboxPatch((0.5, 1.5), 2.8, 3.0,
                               boxstyle="round,pad=0.08",
                               facecolor="#EBF5FB",
                               edgecolor="#2980B9", lw=2)
    ax.add_patch(user_box)
    ax.text(1.9, 4.15, "User", ha="center", fontsize=11,
            fontweight="bold", color="#2980B9")
    ax.plot([0.6, 3.2], [3.95, 3.95], color="#2980B9", lw=1)
    fields = ["id: UUID (PK)", "name: str", "age: int",
              "gender: str", "zone_id: str"]
    for i, f in enumerate(fields):
        ax.text(1.0, 3.7 - i * 0.4, f, fontsize=7.5,
                fontfamily="monospace", color="#2C3E50")

    # Swipe entity
    swipe_box = FancyBboxPatch((4.0, 1.5), 2.5, 3.0,
                                boxstyle="round,pad=0.08",
                                facecolor="#F5EEF8",
                                edgecolor="#8E44AD", lw=2)
    ax.add_patch(swipe_box)
    ax.text(5.25, 4.15, "Swipe", ha="center", fontsize=11,
            fontweight="bold", color="#8E44AD")
    ax.plot([4.1, 6.4], [3.95, 3.95], color="#8E44AD", lw=1)
    fields = ["swiper_id: UUID (FK)", "swiped_id: UUID (FK)",
              "action: SwipeAction", "timestamp: datetime"]
    for i, f in enumerate(fields):
        ax.text(4.3, 3.7 - i * 0.4, f, fontsize=7.5,
                fontfamily="monospace", color="#2C3E50")

    # Match entity
    match_box = FancyBboxPatch((7.2, 1.5), 2.8, 3.0,
                                boxstyle="round,pad=0.08",
                                facecolor="#EAFAF1",
                                edgecolor="#27AE60", lw=2)
    ax.add_patch(match_box)
    ax.text(8.6, 4.15, "Match", ha="center", fontsize=11,
            fontweight="bold", color="#27AE60")
    ax.plot([7.3, 9.9], [3.95, 3.95], color="#27AE60", lw=1)
    fields = ["user1_id: UUID (FK)", "user2_id: UUID (FK)",
              "timestamp: datetime"]
    for i, f in enumerate(fields):
        ax.text(7.5, 3.7 - i * 0.4, f, fontsize=7.5,
                fontfamily="monospace", color="#2C3E50")

    # SwipeAction enum
    enum_box = FancyBboxPatch((4.2, 0, ), 2.1, 1.0,
                               boxstyle="round,pad=0.08",
                               facecolor="#FEF9E7",
                               edgecolor="#F39C12", lw=1.5)
    ax.add_patch(enum_box)
    ax.text(5.25, 0.7, "SwipeAction", ha="center", fontsize=9,
            fontweight="bold", color="#F39C12")
    ax.text(5.25, 0.35, "LIKE | PASS", ha="center", fontsize=8,
            fontfamily="monospace", color="#2C3E50")

    # Relationship arrows
    # User -> Swipe (swiper)
    ax.annotate("", xy=(4.0, 3.6), xytext=(3.3, 3.6),
                arrowprops=dict(arrowstyle="-|>", color="#7F8C8D", lw=1.5))
    ax.text(3.55, 3.75, "1:N", fontsize=7, color="#7F8C8D", ha="center")
    ax.text(3.55, 3.45, "swiper", fontsize=6.5, color="#7F8C8D",
            ha="center", style="italic")

    # User -> Swipe (swiped)
    ax.annotate("", xy=(4.0, 3.1), xytext=(3.3, 3.1),
                arrowprops=dict(arrowstyle="-|>", color="#7F8C8D", lw=1.5))
    ax.text(3.55, 2.85, "swiped", fontsize=6.5, color="#7F8C8D",
            ha="center", style="italic")

    # Swipe -> Match
    ax.annotate("", xy=(7.2, 3.3), xytext=(6.5, 3.3),
                arrowprops=dict(arrowstyle="-|>", color="#7F8C8D", lw=1.5))
    ax.text(6.8, 3.5, "mutual LIKEs\ncreate Match", fontsize=6.5,
            color="#7F8C8D", ha="center", style="italic")

    # User -> Match
    ax.annotate("", xy=(7.2, 2.5), xytext=(3.3, 2.2),
                arrowprops=dict(arrowstyle="-|>", color="#BDC3C7",
                                lw=1.2, connectionstyle="arc3,rad=-0.15"))
    ax.text(5.2, 1.8, "user1 / user2  (1:N)", fontsize=6.5,
            color="#BDC3C7", ha="center", style="italic")

    # SwipeAction -> Swipe
    ax.annotate("", xy=(5.25, 1.5), xytext=(5.25, 1.0),
                arrowprops=dict(arrowstyle="-|>", color="#F39C12", lw=1))

    return _save_fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# Diagram 6: User Journey / Sequence Flow
# ---------------------------------------------------------------------------

def create_user_journey_diagram() -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(10, 8))
    ax.set_xlim(-0.5, 10.5)
    ax.set_ylim(-1, 9)
    ax.axis("off")
    ax.set_title("End-to-End User Journey – Sequence Flow", fontsize=14,
                 fontweight="bold", pad=15)

    # Participants
    parts = {
        "Alice": (1.2, "#E67E22"),
        "API": (4.0, "#3498DB"),
        "Services": (6.5, "#9B59B6"),
        "Store": (9.0, "#E74C3C"),
    }
    for name, (x, c) in parts.items():
        _draw_box(ax, x - 0.6, 7.8, 1.2, 0.6, name, c, "white", 9)
        ax.plot([x, x], [0, 7.8], color="#BDC3C7", lw=0.8, linestyle="--")

    steps = [
        (7.0, "Alice", "API", "POST /users/  {name, age, ...}", "#27AE60"),
        (6.5, "API", "Services", "validate + create User", "#27AE60"),
        (6.0, "Services", "Store", "add_user(user)", "#27AE60"),
        (5.5, "Store", "Alice", "← 201  {id: uuid-A}", "#27AE60"),

        (4.8, "Alice", "API", "GET /feed?user_id=uuid-A", "#2980B9"),
        (4.3, "API", "Services", "FeedService.generate_feed()", "#2980B9"),
        (3.8, "Services", "Store", "get_all_users + get_swipes", "#2980B9"),
        (3.3, "Store", "Alice", "← 200  {data: [Bob, ...]}", "#2980B9"),

        (2.6, "Alice", "API", "POST /swipe  {A→B, LIKE}", "#8E44AD"),
        (2.1, "API", "Services", "SwipeService.process_swipe()", "#8E44AD"),
        (1.6, "Services", "Store", "add_swipe + find_reverse", "#8E44AD"),
        (1.1, "Store", "Alice", "← 201  {is_match: true/false}", "#8E44AD"),

        (0.4, "Alice", "API", "GET /matches?user_id=uuid-A", "#D35400"),
        (-0.1, "API", "Store", "get_matches_for_user()", "#D35400"),
        (-0.6, "Store", "Alice", "← 200  {data: [Match, ...]}", "#D35400"),
    ]

    for y, src, dst, label, color in steps:
        x1 = parts[src][0]
        x2 = parts[dst][0]
        direction = 1 if x2 > x1 else -1
        ax.annotate(
            "", xy=(x2, y), xytext=(x1, y),
            arrowprops=dict(arrowstyle="-|>", color=color, lw=1.3)
        )
        mx = (x1 + x2) / 2
        ax.text(mx, y + 0.12, label, fontsize=6.5, ha="center",
                color=color, fontweight="bold")

    return _save_fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# DOCX Document Builder
# ---------------------------------------------------------------------------

def build_docx():
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    # ============================= TITLE PAGE ==============================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Tinder-Claude")
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("System Design Document")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Version 1.0  •  February 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.add_page_break()

    # ========================= TABLE OF CONTENTS ===========================
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. System Architecture",
        "    2.1 High-Level Architecture Diagram",
        "    2.2 Technology Stack",
        "    2.3 Project Structure",
        "3. Data Models",
        "    3.1 Entity Relationship Diagram",
        "    3.2 Model Definitions",
        "4. API Design",
        "    4.1 Endpoint Summary",
        "    4.2 Response Envelope",
        "    4.3 Error Handling",
        "5. Data Flow Diagrams",
        "    5.1 User Request Lifecycle",
        "    5.2 Feed Generation Pipeline",
        "    5.3 Match Detection Flow",
        "6. User Journey – Sequence Flow",
        "7. Service Layer Details",
        "    7.1 InMemoryStore",
        "    7.2 FeedService",
        "    7.3 SwipeService",
        "8. Testing Strategy",
        "9. Security Considerations",
        "10. Future Enhancements",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        if not item.startswith("    "):
            p.runs[0].bold = True
    doc.add_page_break()

    # ====================== 1. EXECUTIVE SUMMARY ==========================
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Tinder-Claude is a backend REST API prototype that implements core "
        "dating-app mechanics: profile creation, location-based discovery feeds, "
        "swiping interactions (LIKE/PASS), and mutual match detection. Built with "
        "FastAPI and Python 3.11+, it serves as an educational reference for "
        "developers studying matching systems and geo-spatial filtering logic."
    )
    doc.add_paragraph(
        "The system uses an in-memory singleton store for data persistence, "
        "making it lightweight and easy to test while clearly demonstrating "
        "the architectural patterns that would scale to a production system "
        "backed by a real database."
    )

    key_points = [
        "RESTful API with six endpoints covering the complete user lifecycle",
        "Three-tier feed filtering: zone-based → self-exclusion → seen-state",
        "Bidirectional match detection on mutual LIKE actions",
        "Pydantic v2 models for strict request/response validation",
        "Comprehensive test suite with unit and integration tests",
        "Standardized API response envelope (data, meta, errors)",
    ]
    doc.add_paragraph("Key characteristics:")
    for point in key_points:
        doc.add_paragraph(point, style="List Bullet")
    doc.add_page_break()

    # ====================== 2. SYSTEM ARCHITECTURE =========================
    doc.add_heading("2. System Architecture", level=1)

    doc.add_heading("2.1 High-Level Architecture Diagram", level=2)
    doc.add_paragraph(
        "The system follows a layered architecture with clear separation "
        "of concerns. HTTP requests flow from the client through the FastAPI "
        "application layer, are routed to endpoint handlers, which delegate "
        "business logic to specialized services that interact with the "
        "centralized in-memory store."
    )
    arch_img = create_architecture_diagram()
    doc.add_picture(arch_img, width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("2.2 Technology Stack", level=2)
    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.style = "Light Grid Accent 1"
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tech_table.rows[0].cells
    hdr[0].text = "Layer"
    hdr[1].text = "Technology"
    hdr[2].text = "Purpose"
    for h in hdr:
        for p in h.paragraphs:
            for r in p.runs:
                r.bold = True

    tech_data = [
        ("Language", "Python 3.11+", "Core runtime with modern type hints"),
        ("Framework", "FastAPI ≥0.110", "High-performance async web framework"),
        ("Server", "Uvicorn ≥0.29", "ASGI server for FastAPI"),
        ("Validation", "Pydantic ≥2.6", "Data parsing and validation"),
        ("Storage", "In-Memory (dict/list)", "Prototype data persistence"),
        ("Testing", "pytest ≥8.0 + httpx", "Unit & integration testing"),
        ("Linting", "ruff ≥0.3", "Code formatting and lint checks"),
        ("Package Mgr", "uv", "Fast Python package management"),
    ]
    for layer, tech, purpose in tech_data:
        row = tech_table.add_row().cells
        row[0].text = layer
        row[1].text = tech
        row[2].text = purpose

    doc.add_paragraph()  # spacer

    doc.add_heading("2.3 Project Structure", level=2)
    structure = (
        "tinder-claude/\n"
        "├── app/\n"
        "│   ├── main.py              # FastAPI app entry point, router setup\n"
        "│   ├── models/\n"
        "│   │   └── schemas.py       # Pydantic data models & enums\n"
        "│   ├── services/\n"
        "│   │   ├── store.py         # InMemoryStore singleton\n"
        "│   │   ├── feed_service.py  # Feed generation with 3-tier filter\n"
        "│   │   └── swipe_service.py # Swipe processing & match detection\n"
        "│   └── routes/\n"
        "│       ├── users.py         # POST/GET user endpoints\n"
        "│       ├── feed.py          # GET feed endpoint\n"
        "│       └── swipe.py         # POST swipe, GET matches endpoints\n"
        "├── tests/\n"
        "│   ├── conftest.py          # Shared fixtures (client, users)\n"
        "│   ├── test_api.py          # Integration tests (35+ scenarios)\n"
        "│   ├── test_store.py        # Store unit tests\n"
        "│   ├── test_feed_service.py # Feed service unit tests\n"
        "│   └── test_swipe_service.py# Swipe service unit tests\n"
        "├── pyproject.toml           # Dependencies & tool configuration\n"
        "└── README.md                # User-facing documentation\n"
    )
    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    doc.add_page_break()

    # ====================== 3. DATA MODELS ================================
    doc.add_heading("3. Data Models", level=1)

    doc.add_heading("3.1 Entity Relationship Diagram", level=2)
    doc.add_paragraph(
        "The system uses three core entities: User, Swipe, and Match. Users "
        "generate Swipes; when two users mutually LIKE each other, a Match "
        "record is created. All models use UUID identifiers and UTC timestamps."
    )
    er_img = create_data_model_diagram()
    doc.add_picture(er_img, width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3.2 Model Definitions", level=2)

    # User model table
    doc.add_paragraph("User Model", style="Heading 3")
    user_tbl = doc.add_table(rows=1, cols=4)
    user_tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(["Field", "Type", "Default", "Description"]):
        user_tbl.rows[0].cells[i].text = h
        for r in user_tbl.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
    user_fields = [
        ("id", "UUID", "uuid4()", "Unique identifier"),
        ("name", "str", "required", "Display name"),
        ("age", "int", "required", "User age"),
        ("gender", "str", "required", "Gender identifier"),
        ("zone_id", "str", "required", "Geographic zone for feed filtering"),
    ]
    for field, typ, default, desc in user_fields:
        row = user_tbl.add_row().cells
        row[0].text = field
        row[1].text = typ
        row[2].text = default
        row[3].text = desc

    doc.add_paragraph()

    # Swipe model table
    doc.add_paragraph("Swipe Model", style="Heading 3")
    swipe_tbl = doc.add_table(rows=1, cols=4)
    swipe_tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(["Field", "Type", "Default", "Description"]):
        swipe_tbl.rows[0].cells[i].text = h
        for r in swipe_tbl.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
    swipe_fields = [
        ("swiper_id", "UUID", "required", "ID of the user who swiped"),
        ("swiped_id", "UUID", "required", "ID of the user being swiped on"),
        ("action", "SwipeAction", "required", "LIKE or PASS"),
        ("timestamp", "datetime", "utcnow()", "When the swipe occurred"),
    ]
    for field, typ, default, desc in swipe_fields:
        row = swipe_tbl.add_row().cells
        row[0].text = field
        row[1].text = typ
        row[2].text = default
        row[3].text = desc

    doc.add_paragraph()

    # Match model table
    doc.add_paragraph("Match Model", style="Heading 3")
    match_tbl = doc.add_table(rows=1, cols=4)
    match_tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(["Field", "Type", "Default", "Description"]):
        match_tbl.rows[0].cells[i].text = h
        for r in match_tbl.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
    match_fields = [
        ("user1_id", "UUID", "required", "First user in the match"),
        ("user2_id", "UUID", "required", "Second user in the match"),
        ("timestamp", "datetime", "utcnow()", "When the match was created"),
    ]
    for field, typ, default, desc in match_fields:
        row = match_tbl.add_row().cells
        row[0].text = field
        row[1].text = typ
        row[2].text = default
        row[3].text = desc

    doc.add_page_break()

    # ====================== 4. API DESIGN =================================
    doc.add_heading("4. API Design", level=1)

    doc.add_heading("4.1 Endpoint Summary", level=2)
    api_tbl = doc.add_table(rows=1, cols=5)
    api_tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(["Method", "Endpoint", "Status", "Description",
                            "Error Codes"]):
        api_tbl.rows[0].cells[i].text = h
        for r in api_tbl.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
    endpoints = [
        ("GET", "/", "200", "Health check", "—"),
        ("POST", "/users/", "201", "Create a new user profile", "422"),
        ("GET", "/users/{id}", "200", "Retrieve user by UUID", "404"),
        ("GET", "/feed?user_id=", "200", "Get filtered discovery feed",
         "404"),
        ("POST", "/swipe", "201", "Submit a swipe action", "400, 422"),
        ("GET", "/matches?user_id=", "200", "List matches for a user",
         "404"),
    ]
    for method, endpoint, status, desc, errors in endpoints:
        row = api_tbl.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = status
        row[3].text = desc
        row[4].text = errors

    doc.add_paragraph()

    doc.add_heading("4.2 Response Envelope", level=2)
    doc.add_paragraph(
        "All API responses are wrapped in a standardized ApiResponse envelope "
        "to provide a consistent interface for clients:"
    )
    envelope_code = (
        '{\n'
        '    "data": <primary_payload>,     // Main response data\n'
        '    "meta": {"count": N, ...},     // Metadata (counts, pagination)\n'
        '    "errors": [{"message": "..."}] // Error details (if any)\n'
        '}'
    )
    p = doc.add_paragraph()
    run = p.add_run(envelope_code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    doc.add_heading("4.3 Error Handling", level=2)
    doc.add_paragraph(
        "The API uses standard HTTP status codes with descriptive error messages:"
    )
    error_items = [
        "400 Bad Request — Business logic errors (self-swipe, invalid user references in swipe)",
        "404 Not Found — Requested user or resource does not exist",
        "422 Unprocessable Entity — Request body fails Pydantic validation (missing/invalid fields)",
    ]
    for item in error_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ====================== 5. DATA FLOW DIAGRAMS =========================
    doc.add_heading("5. Data Flow Diagrams", level=1)

    doc.add_heading("5.1 User Request Lifecycle", level=2)
    doc.add_paragraph(
        "Each API request follows a consistent pipeline: validation, business "
        "logic processing, store interaction, and response formatting. The "
        "diagram below shows the data flow for all four primary operations."
    )
    lifecycle_img = create_request_lifecycle_diagram()
    doc.add_picture(lifecycle_img, width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    doc.add_heading("5.2 Feed Generation Pipeline", level=2)
    doc.add_paragraph(
        "The feed generation system applies a three-tier filtering pipeline "
        "to produce a personalized discovery feed. Each filter progressively "
        "narrows the candidate pool:"
    )
    filter_items = [
        "Zone Filter — Only users in the same geographic zone (zone_id match)",
        "Self-Exclusion — Remove the requesting user from results",
        "Seen-State Filter — Remove users already swiped (LIKE or PASS) by the requester",
    ]
    for item in filter_items:
        doc.add_paragraph(item, style="List Bullet")

    feed_img = create_feed_pipeline_diagram()
    doc.add_picture(feed_img, width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    doc.add_heading("5.3 Match Detection Flow", level=2)
    doc.add_paragraph(
        "Match detection occurs during swipe processing. When a LIKE swipe is "
        "recorded, the system checks for a reverse LIKE from the target user. "
        "If both users have liked each other, a Match record is created and "
        "the response indicates a successful match."
    )
    match_img = create_match_detection_diagram()
    doc.add_picture(match_img, width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("Match detection algorithm:")
    algo_code = (
        "def process_swipe(swiper_id, swiped_id, action):\n"
        "    store.add_swipe(Swipe(swiper_id, swiped_id, action))\n"
        "    if action == LIKE:\n"
        "        reverse = store.find_swipe(swiped_id, swiper_id)\n"
        "        if reverse and reverse.action == LIKE:\n"
        "            store.add_match(Match(swiper_id, swiped_id))\n"
        "            return True  # It's a match!\n"
        "    return False"
    )
    p = doc.add_paragraph()
    run = p.add_run(algo_code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    doc.add_page_break()

    # ====================== 6. USER JOURNEY ===============================
    doc.add_heading("6. User Journey – Sequence Flow", level=1)
    doc.add_paragraph(
        "The following sequence diagram illustrates a complete user journey "
        "from profile creation through feed discovery, swiping, and match "
        "retrieval. It shows the interaction between four system components: "
        "the client (Alice), the API layer, the service layer, and the data store."
    )
    journey_img = create_user_journey_diagram()
    doc.add_picture(journey_img, width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("Typical user journey steps:")
    journey_steps = [
        "Create Profile — User registers with name, age, gender, and zone. Receives UUID.",
        "Discover Feed — User requests discovery feed filtered by zone, excluding self and already-swiped profiles.",
        "Swipe Actions — User submits LIKE or PASS decisions. Each swipe is recorded and checked for a mutual match.",
        "View Matches — User retrieves list of all mutual matches for further interaction.",
    ]
    for i, step in enumerate(journey_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_page_break()

    # ====================== 7. SERVICE LAYER ==============================
    doc.add_heading("7. Service Layer Details", level=1)

    doc.add_heading("7.1 InMemoryStore", level=2)
    doc.add_paragraph(
        "The InMemoryStore is a singleton class that serves as the centralized "
        "data persistence layer. It maintains three collections:"
    )
    store_items = [
        "users: dict[UUID, User] — O(1) lookup by user ID",
        "swipes: list[Swipe] — Chronologically ordered swipe records",
        "matches: list[Match] — Chronologically ordered match records",
    ]
    for item in store_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "The singleton pattern ensures all services share the same data "
        "state. A reset() method is provided for test isolation. Key "
        "operations include add_user(), get_user(), get_all_users(), "
        "add_swipe(), get_swipes_by_user(), find_swipe(), add_match(), "
        "and get_matches_for_user()."
    )

    doc.add_heading("7.2 FeedService", level=2)
    doc.add_paragraph(
        "FeedService is responsible for generating personalized discovery feeds. "
        "It receives a user ID and applies the three-tier filter pipeline "
        "(zone → self → seen-state) to produce a list of candidate profiles. "
        "The service raises a ValueError if the requesting user does not exist."
    )
    doc.add_paragraph(
        "Performance characteristics: The current implementation iterates "
        "through all users in the store (O(N)) and builds a set of "
        "already-swiped IDs for O(1) membership checks. This is efficient "
        "for the prototype scale but would require indexing by zone_id "
        "for production workloads."
    )

    doc.add_heading("7.3 SwipeService", level=2)
    doc.add_paragraph(
        "SwipeService handles swipe processing and match detection. It "
        "validates that both users exist, prevents self-swiping, records "
        "the swipe action, and on LIKE actions, checks for a reciprocal "
        "LIKE to detect mutual matches."
    )
    doc.add_paragraph(
        "The match detection is O(N) on the swipes list in the worst case "
        "(via find_swipe). For production, this would benefit from an "
        "indexed lookup structure or database query optimization."
    )

    doc.add_page_break()

    # ====================== 8. TESTING STRATEGY ===========================
    doc.add_heading("8. Testing Strategy", level=1)
    doc.add_paragraph(
        "The project includes a comprehensive test suite organized into "
        "unit tests for individual services and integration tests for the "
        "full API."
    )

    test_tbl = doc.add_table(rows=1, cols=4)
    test_tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(["Test File", "Type", "Coverage",
                            "Key Scenarios"]):
        test_tbl.rows[0].cells[i].text = h
        for r in test_tbl.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
    tests = [
        ("test_store.py", "Unit",
         "InMemoryStore", "Singleton behavior, CRUD, reset"),
        ("test_feed_service.py", "Unit",
         "FeedService", "Zone filtering, self-exclusion, seen-state"),
        ("test_swipe_service.py", "Unit",
         "SwipeService", "Match detection, self-swipe prevention"),
        ("test_api.py", "Integration",
         "All endpoints", "35+ scenarios, error cases, full flows"),
    ]
    for f, t, cov, scen in tests:
        row = test_tbl.add_row().cells
        row[0].text = f
        row[1].text = t
        row[2].text = cov
        row[3].text = scen

    doc.add_paragraph()
    doc.add_paragraph(
        "Test fixtures in conftest.py provide an auto-use store reset "
        "fixture to ensure test isolation, a shared HTTP test client, "
        "and pre-configured sample users across multiple zones."
    )

    doc.add_page_break()

    # ====================== 9. SECURITY ===================================
    doc.add_heading("9. Security Considerations", level=1)
    doc.add_paragraph(
        "As a prototype, the current implementation does not include "
        "authentication or authorization. The following areas should be "
        "addressed before any production deployment:"
    )
    security_items = [
        "Authentication — Implement JWT or OAuth2 token-based authentication",
        "Authorization — Ensure users can only perform actions on their own behalf",
        "Rate Limiting — Protect against abuse with per-user request throttling",
        "Input Sanitization — While Pydantic handles type validation, additional "
        "business rule validation may be needed",
        "HTTPS — All production traffic should be encrypted in transit",
        "Data Persistence — Replace in-memory store with a proper database "
        "with encryption at rest",
        "Audit Logging — Track sensitive operations for security monitoring",
    ]
    for item in security_items:
        doc.add_paragraph(item, style="List Bullet")

    # ====================== 10. FUTURE ENHANCEMENTS =======================
    doc.add_heading("10. Future Enhancements", level=1)
    doc.add_paragraph(
        "The following enhancements would bring the system closer to "
        "production readiness:"
    )
    future_items = [
        "Database Integration — PostgreSQL with SQLAlchemy ORM for persistent storage",
        "Real-time Notifications — WebSocket support for instant match notifications",
        "Geolocation — GPS coordinate-based distance filtering instead of zone strings",
        "Recommendation Engine — ML-based profile suggestions beyond simple zone filtering",
        "Media Support — Profile photo upload and storage (S3/cloud storage integration)",
        "Messaging — In-app chat system for matched users",
        "Pagination — Cursor-based pagination for feed and match endpoints",
        "Caching — Redis layer for frequently accessed data (feeds, user profiles)",
        "Containerization — Docker support with docker-compose for deployment",
        "CI/CD — Automated testing and deployment pipeline",
    ]
    for item in future_items:
        doc.add_paragraph(item, style="List Bullet")

    # ======================== SAVE ========================================
    output_path = os.path.join(OUTPUT_DIR, "design_document.docx")
    doc.save(output_path)
    print(f"Design document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    build_docx()
